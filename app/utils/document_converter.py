# app/utils/document_converter.py

import pdfplumber
from docx import Document
import os
import tempfile
from fastapi import UploadFile, HTTPException


async def convert_to_markdown(file: UploadFile) -> str:
    try:
        content = await file.read()
        extension = os.path.splitext(file.filename)[1].lower()

        with tempfile.NamedTemporaryFile(delete=False, suffix=extension) as temp_file:
            temp_file.write(content)
            temp_file_path = temp_file.name

        if extension == '.pdf':
            result = await convert_pdf_to_markdown(temp_file_path)
        elif extension in ['.doc', '.docx']:
            result = await convert_doc_to_markdown(temp_file_path)
        else:
            os.unlink(temp_file_path)
            raise HTTPException(status_code=400, detail=f"Unsupported file format: {extension}")

        os.unlink(temp_file_path)
        return result
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        raise HTTPException(status_code=500, detail=f"Error converting file: {str(e)}\n\nDetails:\n{error_details}")


async def convert_pdf_to_markdown(file_path: str) -> str:
    try:
        markdown = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                tables = page.extract_tables()

                # Convert text to markdown (simple heuristic for headers)
                lines = text.split('\n')
                for line in lines:
                    if len(line.strip()) > 0:
                        if len(line) < 50 and line.isupper():
                            markdown += f"# {line}\n\n"
                        else:
                            markdown += f"{line}\n\n"

                # Convert tables to markdown
                for table in tables:
                    if table:  # Check if table is not empty
                        # Convert all cell values to strings, replacing None with empty string
                        table = [[str(cell) if cell is not None else '' for cell in row] for row in table]
                        markdown += "| " + " | ".join(table[0]) + " |\n"
                        markdown += "| " + " | ".join(["---" for _ in table[0]]) + " |\n"
                        for row in table[1:]:
                            markdown += "| " + " | ".join(row) + " |\n"
                        markdown += "\n"

        return markdown
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error converting PDF: {str(e)}")


async def convert_doc_to_markdown(file_path: str) -> str:
    try:
        doc = Document(file_path)
        markdown = ""

        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                level = int(para.style.name[-1])
                markdown += f"{'#' * level} {para.text}\n\n"
            elif para.style.name == 'List Paragraph':
                markdown += f"- {para.text}\n"
            else:
                markdown += f"{para.text}\n\n"

        for table in doc.tables:
            for row in table.rows:
                markdown += "| " + " | ".join(cell.text for cell in row.cells) + " |\n"
            markdown += "\n"

        return markdown
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error converting DOC/DOCX: {str(e)}")