from typing import Optional, Dict, Any
import os
from pathlib import Path
import tempfile
import mimetypes
from fastapi import HTTPException, UploadFile
import textract  # Main extraction library
import pdfplumber  # Better PDF handling
from bs4 import BeautifulSoup  # HTML handling
import docx  # Specialized DOCX handling
from datetime import datetime


class DocumentExtractor:
    """Universal document extraction class that handles multiple document formats."""

    SUPPORTED_MIMETYPES = {
        # Text documents
        'text/plain': 'txt',
        'text/html': 'html',
        'text/csv': 'csv',
        # Microsoft Office
        'application/msword': 'doc',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'application/vnd.ms-excel': 'xls',
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
        'application/vnd.ms-powerpoint': 'ppt',
        'application/vnd.openxmlformats-officedocument.presentationml.presentation': 'pptx',
        # PDF
        'application/pdf': 'pdf',
        # Email
        'message/rfc822': 'eml',
        'application/vnd.ms-outlook': 'msg',
        # OpenOffice/LibreOffice
        'application/vnd.oasis.opendocument.text': 'odt',
        'application/vnd.oasis.opendocument.spreadsheet': 'ods',
        'application/vnd.oasis.opendocument.presentation': 'odp',
        # Images (with OCR)
        'image/jpeg': 'jpg',
        'image/png': 'png',
        'image/tiff': 'tiff'
    }

    def __init__(self, ocr_enabled: bool = True, email_enabled: bool = True):
        """Initialize the document extractor.

        Args:
            ocr_enabled (bool): Whether to enable OCR for images and scanned PDFs
            email_enabled (bool): Whether to enable email parsing
        """
        self.ocr_enabled = ocr_enabled
        self.email_enabled = email_enabled
        self._initialize_parsers()

    def _initialize_parsers(self):
        """Initialize parsers based on enabled features"""
        if self.email_enabled:
            try:
                import eml_parser
                self.eml_parser = eml_parser
            except ImportError:
                self.email_enabled = False
                print("Warning: Email parsing disabled due to missing dependencies")

    async def extract_text(self, file: UploadFile) -> Dict[str, Any]:
        """Extract text and metadata from a document.

        Args:
            file (UploadFile): The uploaded file to process

        Returns:
            Dict containing extracted text and metadata
        """
        try:
            # Detect mimetype
            mime_type = file.content_type or mimetypes.guess_type(file.filename)[0]
            if not mime_type or mime_type not in self.SUPPORTED_MIMETYPES:
                raise HTTPException(status_code=400, detail=f"Unsupported file type: {mime_type}")

            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False,
                                             suffix=f".{self.SUPPORTED_MIMETYPES[mime_type]}") as temp_file:
                content = await file.read()
                temp_file.write(content)
                temp_path = temp_file.name

            try:
                # Extract based on file type
                if mime_type == 'application/pdf':
                    text, metadata = self._extract_from_pdf(temp_path)
                elif mime_type in ['application/msword',
                                   'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                    text, metadata = self._extract_from_doc(temp_path)
                elif mime_type == 'text/html':
                    text, metadata = self._extract_from_html(content)
                elif mime_type in ['application/vnd.ms-excel',
                                   'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']:
                    text, metadata = self._extract_from_excel(temp_path)
                elif mime_type in ['message/rfc822', 'application/vnd.ms-outlook']:
                    text, metadata = self._extract_from_email(temp_path)
                elif mime_type.startswith('image/'):
                    text, metadata = self._extract_from_image(temp_path)
                else:
                    # Fallback to textract for other formats
                    text = textract.process(temp_path).decode('utf-8')
                    metadata = self._get_basic_metadata(temp_path)

                return {
                    "filename": file.filename,
                    "mime_type": mime_type,
                    "text": text,
                    "metadata": metadata,
                    "extraction_timestamp": datetime.now().isoformat()
                }

            finally:
                # Clean up temp file
                os.unlink(temp_path)

        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error extracting text: {str(e)}")

    def _extract_from_pdf(self, filepath: str) -> tuple[str, Dict[str, Any]]:
        """Extract text and metadata from PDF files."""
        text = ""
        metadata = {}

        with pdfplumber.open(filepath) as pdf:
            metadata = pdf.metadata
            for page in pdf.pages:
                text += page.extract_text() or ""
                # Extract tables if present
                tables = page.extract_tables()
                if tables:
                    text += "\n\nTables:\n"
                    for table in tables:
                        text += "\n".join([", ".join([str(cell) for cell in row if cell]) for row in table]) + "\n"

        return text.strip(), metadata

    def _extract_from_doc(self, filepath: str) -> tuple[str, Dict[str, Any]]:
        """Extract text and metadata from DOC/DOCX files."""
        try:
            doc = docx.Document(filepath)

            # Extract text from paragraphs
            text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs])

            # Extract tables
            tables_text = []
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                tables_text.append("\n".join(table_text))

            if tables_text:
                text += "\n\nTables:\n" + "\n\n".join(tables_text)

            # Extract metadata
            metadata = {
                "core_properties": {
                    "author": doc.core_properties.author,
                    "created": doc.core_properties.created.isoformat() if doc.core_properties.created else None,
                    "modified": doc.core_properties.modified.isoformat() if doc.core_properties.modified else None,
                    "title": doc.core_properties.title,
                    "subject": doc.core_properties.subject
                }
            }

            return text.strip(), metadata

        except Exception:
            # Fallback to textract for older DOC files
            text = textract.process(filepath).decode('utf-8')
            return text.strip(), self._get_basic_metadata(filepath)

    def _extract_from_html(self, content: bytes) -> tuple[str, Dict[str, Any]]:
        """Extract text and metadata from HTML content."""
        soup = BeautifulSoup(content, 'html.parser')

        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()

        # Extract metadata from meta tags
        metadata = {
            "title": soup.title.string if soup.title else None,
            "meta": {
                meta.get('name', meta.get('property')): meta.get('content')
                for meta in soup.find_all('meta')
                if meta.get('name') or meta.get('property')
            }
        }

        # Get text content
        text = soup.get_text(separator='\n', strip=True)
        return text, metadata

    def _extract_from_excel(self, filepath: str) -> tuple[str, Dict[str, Any]]:
        """Extract text and metadata from Excel files."""
        try:
            import pandas as pd
            # Read all sheets
            excel_file = pd.ExcelFile(filepath)
            sheets_data = []

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(filepath, sheet_name=sheet_name)
                sheets_data.append(f"Sheet: {sheet_name}\n{df.to_string()}")

            text = "\n\n".join(sheets_data)
            metadata = {
                "sheets": excel_file.sheet_names,
                "file_properties": excel_file.book.properties.__dict__
            }

            return text, metadata

        except Exception as e:
            # Fallback to textract
            text = textract.process(filepath).decode('utf-8')
            return text, self._get_basic_metadata(filepath)

    def _extract_from_email(self, filepath: str) -> tuple[str, Dict[str, Any]]:
        """Extract text and metadata from email files."""
        if not self.email_enabled:
            return "Email parsing not enabled", {"error": "Email parsing dependencies not available"}

        try:
            with open(filepath, 'rb') as fhdl:
                raw_email = fhdl.read()

            ep = self.eml_parser.EmlParser()
            parsed_eml = ep.decode_email_bytes(raw_email)

            # Extract text content
            text_parts = []
            if 'body' in parsed_eml:
                if 'plain' in parsed_eml['body']:
                    text_parts.append(parsed_eml['body']['plain'])
                if 'html' in parsed_eml['body']:
                    # Convert HTML to text
                    from bs4 import BeautifulSoup
                    soup = BeautifulSoup(parsed_eml['body']['html'], 'html.parser')
                    text_parts.append(soup.get_text())

            text = "\n\n".join(text_parts)

            # Extract metadata
            metadata = {
                "subject": parsed_eml.get('subject'),
                "from": parsed_eml.get('from'),
                "to": parsed_eml.get('to'),
                "date": parsed_eml.get('date'),
                "attachments": parsed_eml.get('attachments', [])
            }

            return text, metadata

        except Exception as e:
            # Fallback to basic text extraction
            import textract
            text = textract.process(filepath).decode('utf-8')
            return text, self._get_basic_metadata(filepath)

    def _extract_from_image(self, filepath: str) -> tuple[str, Dict[str, Any]]:
        """Extract text from images using OCR if enabled."""
        if not self.ocr_enabled:
            raise HTTPException(status_code=400, detail="OCR is not enabled for image processing")

        try:
            import pytesseract
            from PIL import Image

            # Open image and perform OCR
            image = Image.open(filepath)
            text = pytesseract.image_to_string(image)

            # Extract image metadata
            metadata = {
                "format": image.format,
                "size": image.size,
                "mode": image.mode,
                "info": image.info
            }

            return text.strip(), metadata

        except ImportError:
            raise HTTPException(status_code=500, detail="OCR support (pytesseract) is not installed")
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error performing OCR: {str(e)}")

    def _get_basic_metadata(self, filepath: str) -> Dict[str, Any]:
        """Get basic file metadata."""
        path = Path(filepath)
        return {
            "size": path.stat().st_size,
            "created": datetime.fromtimestamp(path.stat().st_ctime).isoformat(),
            "modified": datetime.fromtimestamp(path.stat().st_mtime).isoformat(),
            "extension": path.suffix
        }